import argparse
import os
import re
import sys
from typing import List, Dict, Optional
from datetime import datetime

import pandas as pd

# Ключевые слова, характерные для описаний отказов, а не для названий функций
FAILURE_KEYWORDS = [
    "отказ", "потеря", "нарушен", "сбой", "неисправ",
    "failure", "loss", "malfunction", "fault", "defect",
    "hazard", "condition", "failure mode", "failure condition",
    "effect of failure", "failure effect",
]


def archive_old(path: str):
    """
    Если файл path существует, переносит его в подкаталог 'Архив'
    с добавлением штампа даты/времени к имени.
    """
    if not os.path.isfile(path):
        return

    folder = os.path.dirname(path)
    archive_dir = os.path.join(folder, "Архив")
    os.makedirs(archive_dir, exist_ok=True)

    base = os.path.basename(path)
    name, ext = os.path.splitext(base)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    new_name = f"{name}_{ts}{ext}"
    new_path = os.path.join(archive_dir, new_name)

    try:
        os.replace(path, new_path)
    except OSError as e:
        # В GUI это уйдёт в лог консоли, не убьёт процесс
        print(f"Не удалось переместить старый файл '{path}': {e}")


def is_failure_like(text: str) -> bool:
    """Грубая проверка: похоже ли название на описание отказа/сценария, а не на функцию."""
    t = text.lower()
    for kw in FAILURE_KEYWORDS:
        if kw in t:
            return True
    return False


def normalize_code(raw: str) -> str:
    """
    Приводит код к виду с точками:
      F1-2_3        -> F1.2.3
      Ф21-20_01     -> Ф21.20.01
      Ф1.           -> Ф1
      Ф 1 . 1 . 1 . -> Ф1.1.1
    """
    raw = raw.strip()
    if not raw:
        return raw

    head = raw[0]
    tail = raw[1:]

    # убираем пробелы внутри кода
    tail = tail.replace(" ", "")

    # приводим разделители к точкам
    tail = tail.replace("_", ".").replace("-", ".")

    # убираем двойные точки
    while ".." in tail:
        tail = tail.replace("..", ".")

    # убираем точки по краям
    tail = tail.strip(".")

    return head + tail


def sort_key_for_code(code: str):
    """Ключ для осмысленной сортировки: F1.2.10 -> [1, 2, 10]."""
    if not code:
        return []

    body = code[1:]
    parts = body.split(".")
    key = []
    for p in parts:
        if p.isdigit():
            key.append(int(p))
        else:
            key.append(p)
    return key


def get_depth(code: str) -> int:
    """
    Глубина кода:
      F1         -> 1
      F1.2       -> 2
      Ф21.20.01  -> 3
    """
    if not code or len(code) < 2:
        return 0
    body = code[1:]
    if not body:
        return 0
    return len(body.split("."))


def _split_numeric_parts(code: str) -> List[str]:
    """
    Внутренний хелпер: возвращает список числовых частей без первой буквы.
    Для 'Ф21.20.01' -> ['21', '20', '01'].
    """
    if not code or len(code) < 2:
        return []
    body = code[1:]
    parts = [p for p in body.split(".") if p]
    return parts


def is_fi_code(code: str) -> bool:
    """
    Эвристика: код относится к функциям ФИ.

    Предположение:
      - первая группа после буквы — одна цифра (Ф0, Ф1, F2, Ф3.1.2 и т.п.);
      - глубина обычно не больше 4 (Ф1.2.3.4).
    """
    parts = _split_numeric_parts(code)
    if not parts:
        return False

    depth = len(parts)
    first_len = len(parts[0])

    # Одна группа: Ф0, Ф1, F3, Ф9 и т.п. — ФИ верхнего уровня
    if depth == 1:
        return True

    # Если первая группа из одной цифры (включая 0) и глубина до 4 — считаем это ФИ
    if first_len == 1 and depth <= 4:
        return True

    return False


def is_fs_code(code: str) -> bool:
    """
    Эвристика: код относится к функциям систем (ФС).

    Предположение:
      - первая группа после буквы — минимум две цифры (Ф21, Ф23);
      - код типа Ф21.20.01, Ф22.30.10 и т.п.
    """
    parts = _split_numeric_parts(code)
    if not parts:
        return False

    depth = len(parts)
    first_len = len(parts[0])

    if first_len >= 2 and depth >= 1:
        return True

    return False


def _normalize_letter_for_filter(ch: str) -> str:
    """
    Нормализация буквы для фильтра:
      - приводим только к верхнему регистру;
      - не склеиваем латинскую F и кириллическую Ф.
    """
    if not ch:
        return ""
    return ch.upper()


def strip_leading_code_tokens(code: str, name: str) -> str:
    """
    Убирает в начале имени 1–2 «кодовидных» токена,
    если они совпадают с кодом функции или его родителем.

    Пример:
      code = 'F1.1.1'
      name = 'F1-1 Разделение ...' -> 'Разделение ...'
      name = 'F1.1.1 F1.1 Разделение ...' -> 'Разделение ...'
    """
    if not code or not name:
        return name

    norm_code = normalize_code(code)

    parent_guess = ""
    parts = norm_code.split(".")
    if len(parts) > 1:
        parent_guess = ".".join(parts[:-1])

    tokens = name.split()
    if not tokens:
        return name

    cut = 0
    # проверяем первые 2 токена максимум
    for i in range(min(2, len(tokens))):
        tok = tokens[i]
        if not tok or tok[0] not in ("F", "Ф"):
            break

        try:
            tok_norm = normalize_code(tok)
        except Exception:
            break

        if tok_norm == norm_code or (parent_guess and tok_norm == parent_guess):
            cut += 1
        else:
            break

    if cut:
        return " ".join(tokens[cut:]).strip()

    return name


def extract_functions_from_text(
    text: str,
    max_depth: int,
    mode: str,
    code_letter: Optional[str] = None,
    code_prefix: Optional[str] = None,
) -> List[Dict[str, str]]:
    """
    Ищет функции в тексте.

    Форматы:
      F1 Название
      Ф1.2 Название
      F1-2-3 Название
      F1_2_3 Название
      Ф21.20.01 Название

    Условия:
      - код должен стоять в начале строки;
      - строки, похожие на описания отказов, отбрасываются;
      - глубина кода ограничена max_depth (0 = без ограничений);
      - mode:
          * 'fi' — функции ФИ;
          * 'fs' — функции систем (ФС).
      - code_letter: фильтр по букве кода (F / Ф). Если пусто — не фильтруем по букве.
      - code_prefix: числовой порог по первой группе:
          "1"  -> берем коды, где первая группа >= 1  (1,2,3,10,...)
          "20" -> берем коды, где первая группа >= 20 (20,21,99,...)
    """
    # Разрешаем пробелы внутри кода и лишнюю точку в конце
    pattern = re.compile(
        r'^\s*([ФF]\s*\d+(?:\s*[.\-_]\s*\d+){0,4}\.?)\s+(.+)$',
        re.UNICODE | re.MULTILINE,
    )

    # нормализуем фильтры
    if code_letter:
        code_letter = _normalize_letter_for_filter(code_letter[0])

    prefix_int: Optional[int] = None
    if code_prefix:
        code_prefix = code_prefix.strip()
        if code_prefix and code_prefix.isdigit():
            prefix_int = int(code_prefix)
        else:
            prefix_int = None  # ввели мусор — игнорируем фильтр

    results: List[Dict[str, str]] = []

    for match in pattern.finditer(text):
        raw_code = match.group(1)
        raw_name = match.group(2).strip()

        # убираем ведущие маркеры списков: "-", "—" и лишние пробелы
        name = re.sub(r'^[\-\–—\s]+', "", raw_name)
        name = name.strip()
        if not name:
            continue

        # фильтрация строк, похожих на описания отказов
        if is_failure_like(name):
            continue

        code = normalize_code(raw_code)
        if not code:
            continue

        # --- фильтр по букве ---
        if code_letter:
            head = code[0]
            head_norm = _normalize_letter_for_filter(head)
            if head_norm != code_letter:
                continue

        # --- фильтр по первой числовой группе: >= prefix_int ---
        if prefix_int is not None:
            parts = _split_numeric_parts(code)
            if not parts:
                continue
            first_group = parts[0]
            if not first_group.isdigit():
                continue
            if int(first_group) < prefix_int:
                continue

        # фильтрация по типу кода (ФИ / ФС)
        if mode == "fi" and not is_fi_code(code):
            continue
        if mode == "fs" and not is_fs_code(code):
            continue

        # ограничение глубины кода
        if max_depth > 0 and get_depth(code) > max_depth:
            continue

        # убираем дублирующиеся коды в начале имени
        name = strip_leading_code_tokens(code, name)

        results.append(
            {
                "Func_LCN": code,
                "Name": name,
            }
        )

    return results


def infer_hierarchy(functions: List[Dict[str, str]], mode: str) -> List[Dict[str, str]]:
    """
    Определяет иерархию функций по коду.

    Для режима "fi":
      F1.2.3 -> родитель F1.2 (если он есть)

    Для режима "fs":
      иерархия не строится, Parent_LCN всегда пустой.
    """
    if mode == "fs":
        for f in functions:
            f["Parent_LCN"] = ""
        return functions

    functions_sorted = sorted(functions, key=lambda f: sort_key_for_code(f["Func_LCN"]))
    lookup = {f["Func_LCN"]: f for f in functions_sorted}

    for f in functions_sorted:
        parts = f["Func_LCN"].split(".")
        if len(parts) > 1:
            parent_code = ".".join(parts[:-1])
        else:
            parent_code = ""

        if parent_code and parent_code in lookup:
            f["Parent_LCN"] = parent_code
        else:
            f["Parent_LCN"] = ""

    return functions_sorted


def consolidate_functions(functions: List[Dict[str, str]]) -> List[Dict[str, str]]:
    """Схлопывает дубликаты по коду."""
    by_code: Dict[str, Dict[str, str]] = {}
    extras: Dict[str, List[str]] = {}

    functions_sorted = sorted(functions, key=lambda f: sort_key_for_code(f["Func_LCN"]))

    for f in functions_sorted:
        code = f["Func_LCN"]
        name = f.get("Name", "").strip()
        parent = f.get("Parent_LCN", "")

        if code not in by_code:
            by_code[code] = {
                "Func_LCN": code,
                "Parent_LCN": parent,
                "Name": name,
            }
        else:
            if name:
                extras.setdefault(code, []).append(name)

    result: List[Dict[str, str]] = []

    for code in sorted(by_code.keys(), key=sort_key_for_code):
        base = by_code[code]
        desc_list = extras.get(code, [])
        desc_text = ""
        if desc_list:
            desc_text = "\n".join(sorted(set(d for d in desc_list if d)))

        result.append(
            {
                "Func_LCN": base["Func_LCN"],
                "Parent_LCN": base["Parent_LCN"],
                "Name": base["Name"],
                "Description": desc_text,
            }
        )

    return result


def read_file_content(path: str) -> str:
    """
    Универсальное чтение текста из:
      .txt, .docx, .pdf, .xlsx, .xls

    Для Excel читаем все листы книги и превращаем в текст.
    """
    ext = os.path.splitext(path)[1].lower()
    text = ""

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)

            chunks = []

            # параграфы
            for p in doc.paragraphs:
                t = p.text
                if t:
                    chunks.append(t)

            # таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text
                        if t:
                            chunks.append(t)

            text = "\n".join(chunks)
        except Exception as e:
            print("Ошибка чтения DOCX:", e)

    elif ext == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            pages = []
            for page in doc:
                pages.append(page.get_text("text"))
            text = "\n".join(pages)
        except Exception as e:
            print("Ошибка чтения PDF:", e)

    elif ext in [".xlsx", ".xls"]:
        try:
            all_sheets = pd.read_excel(path, sheet_name=None, header=None)
            lines = []

            for sheet_name, df in all_sheets.items():
                for _, row in df.iterrows():
                    vals = []
                    for v in row.tolist():
                        s = str(v).strip()
                        if s and s.lower() != "nan":
                            vals.append(s)
                    if vals:
                        lines.append(" ".join(vals))

            text = "\n".join(lines)
        except Exception as e:
            print("Ошибка чтения Excel:", e)

    else:
        print("Неподдерживаемое расширение:", ext)

    return text


def export_to_excel(
    functions: List[Dict[str, str]],
    fi_name: str,
    output_file: str = "Functions.xlsx",
):
    """Сохраняет функции в Excel в формате шаблона Pragmatica."""
    if not functions:
        print("Не найдено ни одной функции в документе.")
        # важное изменение: сообщаем об ошибке вызывающему коду
        sys.exit(1)

    df = pd.DataFrame(functions)

    df.insert(0, "FI_Обозначение", fi_name)
    if "Products_List" not in df.columns:
        df["Products_List"] = ""

    for col in ["Parent_LCN", "Description", "Products_List"]:
        if col not in df.columns:
            df[col] = ""

    df = df[
        [
            "FI_Обозначение",
            "Func_LCN",
            "Parent_LCN",
            "Name",
            "Description",
            "Products_List",
        ]
    ]

    output_path = os.path.abspath(output_file)

    # Перед записью нового файла утащим старый в Архив
    archive_old(output_path)

    df.to_excel(output_path, index=False, sheet_name="Функции")

    print("Сохранено {} функций в файл: {}".format(len(df), output_path))


def main():
    parser = argparse.ArgumentParser(
        description="Парсер функций из документа в Excel (шаблон Pragmatica)"
    )
    parser.add_argument(
        "input_file",
        help="Путь к файлу (.docx, .pdf, .txt, .xlsx, .xls)",
    )
    parser.add_argument(
        "--fi",
        required=True,
        help="Обозначение функционального экземпляра (ФИ)",
    )
    parser.add_argument(
        "--out",
        default="Functions.xlsx",
        help="Имя выходного Excel файла (по умолчанию Functions.xlsx)",
    )
    parser.add_argument(
        "--max-depth",
        type=int,
        default=0,
        help="Максимальное число уровней кода функции (F1.2.3 = 3). 0 = без ограничений.",
    )
    parser.add_argument(
        "--mode",
        choices=["fi", "fs"],
        default="fi",
        help="Режим парсинга: 'fi' – функции ФИ (иерархия по коду), "
             "'fs' – функции систем (плоский список).",
    )
    parser.add_argument(
        "--code-letter",
        default="",
        help="Буква кода (F или Ф). Если не задана — не фильтруем по букве.",
    )
    parser.add_argument(
        "--code-prefix",
        default="",
        help="Порог по первой числовой группе (например '1', '20'). "
             "Берем только коды, у которых первая группа >= этого значения.",
    )

    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print("Ошибка: файл '{}' не найден.".format(args.input_file))
        sys.exit(1)

    print("Обработка файла:", args.input_file)
    print("Обозначение ФИ:", args.fi)
    print("Режим:", args.mode)
    print("Выходной файл:", args.out)
    print("Максимальная глубина кода:", args.max_depth)
    if args.code_letter:
        print("Буква кода:", args.code_letter)
    if args.code_prefix:
        print("Порог первой группы:", args.code_prefix)

    text = read_file_content(args.input_file)
    if not text.strip():
        print("Не удалось извлечь текст из файла.")
        sys.exit(1)

    functions = extract_functions_from_text(
        text,
        max_depth=args.max_depth,
        mode=args.mode,
        code_letter=args.code_letter or None,
        code_prefix=args.code_prefix or None,
    )
    if not functions:
        print("Не найдено ни одной строки с функцией.")
        sys.exit(1)

    functions = infer_hierarchy(functions, mode=args.mode)
    functions = consolidate_functions(functions)
    export_to_excel(functions, args.fi, args.out)


if __name__ == "__main__":
    main()